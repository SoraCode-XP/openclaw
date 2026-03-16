#!/usr/bin/env python3
"""
Word Document Filler
使用个人信息 JSON 数据填充 Word 文档模板
"""

import argparse
import json
import re
import sys
from pathlib import Path
from docx import Document


def load_person_data(json_path):
    """加载个人信息 JSON 文件"""
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"错误：找不到个人信息文件：{json_path}", file=sys.stderr)
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"错误：JSON 格式错误：{e}", file=sys.stderr)
        sys.exit(1)


def get_nested_value(data, path):
    """
    从嵌套的字典中获取值
    例如：get_nested_value(data, 'basic.name') 返回 data['basic']['name']
    支持数组索引：education[0].university
    """
    # 处理数组索引
    array_pattern = re.compile(r'(\w+)\[(\d+)\]')
    
    parts = path.split('.')
    value = data
    
    try:
        for part in parts:
            # 检查是否有数组索引
            match = array_pattern.match(part)
            if match:
                key = match.group(1)
                index = int(match.group(2))
                value = value[key][index]
            else:
                value = value[part]
        return value
    except (KeyError, TypeError, IndexError):
        return None


def format_education(education_list):
    """格式化教育背景为多行文本"""
    if not education_list:
        return "无教育背景"
    
    lines = []
    for edu in education_list:
        start_date = edu.get('startDate', '')
        end_date = edu.get('endDate', '')
        university = edu.get('university', '')
        major = edu.get('major', '')
        degree = edu.get('degree', '')
        gpa = edu.get('gpa', '')
        
        # 格式：2016.09 - 2020.06  XX大学  计算机科学  本科  GPA: 3.6/4.0
        parts = [
            f"{start_date} - {end_date}",
            university,
            major,
            degree
        ]
        if gpa:
            parts.append(f"GPA: {gpa}")
        
        lines.append("  ".join(parts))
    
    return '\n'.join(lines)


def format_experience(experience_list):
    """格式化工作经历为多行文本"""
    if not experience_list:
        return "无工作经历"
    
    lines = []
    for exp in experience_list:
        start_date = exp.get('startDate', '')
        end_date = exp.get('endDate', '')
        company = exp.get('company', '')
        position = exp.get('position', '')
        description = exp.get('description', '')
        
        # 格式：2023.07 - 至今  XX科技有限公司  后端开发工程师
        header = f"{start_date} - {end_date}  {company}  {position}"
        lines.append(header)
        
        # 工作描述
        if description:
            lines.append(description)
        
        # 添加空行（除了最后一条）
        if exp != experience_list[-1]:
            lines.append("")
    
    return '\n'.join(lines)


def format_certificates(certificates_list):
    """格式化证书列表"""
    if not certificates_list:
        return "无证书"
    return "、".join(certificates_list)


def create_field_mapping():
    """创建字段名到 JSON 路径的映射（扩展版）"""
    return {
        # 基本信息
        '姓名': 'basic.name',
        '英文姓名': 'basic.nameEn',
        '性别': 'basic.gender',
        '年龄': 'basic.age',
        '出生日期': 'basic.birthDate',
        '身份证号': 'basic.idCard',
        '联系电话': 'basic.phone',
        '电子邮箱': 'basic.email',
        '求职意向': 'basic.jobIntention',
        
        # 地址信息
        '现居地址': 'address.current',
        '户籍地址': 'address.registered',
        '邮政编码': 'address.postalCode',
        
        # 教育背景（支持第一条教育记录）
        '学历': 'education[0].degree',
        '毕业院校': 'education[0].university',
        '专业': 'education[0].major',
        '入学时间': 'education[0].startDate',
        '毕业时间': 'education[0].endDate',
        'GPA': 'education[0].gpa',
        
        # 本科（第一条教育记录）
        '本科院校': 'education[0].university',
        '本科专业': 'education[0].major',
        '本科GPA': 'education[0].gpa',
        
        # 硕士（第二条教育记录）
        '硕士院校': 'education[1].university',
        '硕士专业': 'education[1].major',
        '硕士GPA': 'education[1].gpa',
        
        # 工作经历（第一份工作）
        '公司名称': 'experience[0].company',
        '职位': 'experience[0].position',
        '入职时间': 'experience[0].startDate',
        '离职时间': 'experience[0].endDate',
        '工作描述': 'experience[0].description',
        
        # 技能与证书
        '技能证书': 'skills.certificates',
        '专业技能': 'skills.professional',
        '技能掌握': 'skills.professional',
        
        # 自我评价
        '自我评价': 'selfEvaluation',
        '个人评价': 'selfEvaluation',
    }


def replace_placeholders(doc, person_data):
    """
    替换文档中的占位符 ${字段名} 或 {{字段名}}
    返回未找到的字段列表
    """
    field_mapping = create_field_mapping()
    missing_fields = set()
    
    # 正则表达式匹配 ${字段名} 或 {{字段名}}
    pattern = re.compile(r'[\$\{]{1,2}([^}]+)\}{1,2}')
    
    def replace_in_text(text):
        """替换文本中的占位符"""
        matches = pattern.findall(text)
        for match in matches:
            field_name = match
            
            # 尝试所有可能的占位符格式
            possible_placeholders = [
                f'${{{{{field_name}}}}}',  # {{field}}
                f'${{{field_name}}}',       # ${field}
                f'{{{{{field_name}}}}}',    # {{field}} without $
            ]
            
            # 特殊处理复杂字段
            if field_name == '工作经历':
                experience = person_data.get('experience', [])
                value = format_experience(experience)
            elif field_name == '教育背景':
                education = person_data.get('education', [])
                value = format_education(education)
            elif field_name == '技能证书':
                certificates = person_data.get('skills', {}).get('certificates', [])
                value = format_certificates(certificates)
            else:
                # 查找字段映射
                json_path = field_mapping.get(field_name)
                if json_path:
                    value = get_nested_value(person_data, json_path)
                    
                    # 如果是列表，转换为字符串
                    if isinstance(value, list):
                        value = format_certificates(value)
                else:
                    value = None
                
                # 如果找不到值，记录缺失字段
                if value is None:
                    missing_fields.add(field_name)
                    return text  # 保留原文本
            
            # 替换所有可能的占位符格式
            for placeholder in possible_placeholders:
                if placeholder in text:
                    text = text.replace(placeholder, str(value))
        
        return text
    
    # 遍历所有段落
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        new_text = replace_in_text(original_text)
        if new_text != original_text:
            paragraph.text = new_text
    
    # 处理表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    new_text = replace_in_text(original_text)
                    if new_text != original_text:
                        paragraph.text = new_text
    
    return missing_fields
    
    return missing_fields


def main():
    parser = argparse.ArgumentParser(description='使用个人信息 JSON 填充 Word 文档模板')
    parser.add_argument('--template', required=True, help='Word 模板文件路径')
    parser.add_argument('--person-data', required=True, help='个人信息 JSON 文件路径')
    parser.add_argument('--output', required=True, help='输出文件路径')
    
    args = parser.parse_args()
    
    # 加载个人信息
    person_data = load_person_data(args.person_data)
    
    # 加载模板文档
    try:
        doc = Document(args.template)
    except Exception as e:
        print(f"错误：无法打开模板文件 {args.template}: {e}", file=sys.stderr)
        sys.exit(1)
    
    # 替换占位符
    missing_fields = replace_placeholders(doc, person_data)
    
    # 保存输出文档
    try:
        # 确保输出目录存在
        output_path = Path(args.output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        doc.save(args.output)
        print(f"✓ 文档填充完成：{args.output}")
        
        # 如果有缺失字段，提示用户
        if missing_fields:
            print(f"\n⚠ 以下字段在个人信息中未找到，保留为占位符：")
            for field in sorted(missing_fields):
                print(f"  - ${{{field}}}")
            print("\n你可以手动填写这些字段，或更新个人信息 JSON 后重新填充。")
    
    except Exception as e:
        print(f"错误：无法保存文档 {args.output}: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
