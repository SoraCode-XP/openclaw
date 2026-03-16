#!/usr/bin/env python3
"""
Template Generator Script
Extracts text from Word documents and generates templates with placeholders.

Usage:
  python template_generator.py --action extract --input input.docx --output output.txt
  python template_generator.py --action generate --input input.docx --fields fields.json --output template.docx
"""

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Error: python-docx is not installed.")
    print("Please run: pip install python-docx")
    sys.exit(1)


def extract_text(input_docx: str, output_text: str) -> None:
    """
    Extract all text content from a Word document.
    
    Args:
        input_docx: Path to input .docx file
        output_text: Path to save extracted text
    """
    try:
        doc = Document(input_docx)
        
        # Extract text from all paragraphs
        all_text = []
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if text:  # Only include non-empty paragraphs
                all_text.append(f"[第{i}段] {text}")
        
        # Also extract text from tables
        for table_idx, table in enumerate(doc.tables, 1):
            all_text.append(f"\n[表格{table_idx}]")
            for row_idx, row in enumerate(table.rows, 1):
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    all_text.append(f"  行{row_idx}: {' | '.join(row_text)}")
        
        # Save to output file
        output_path = Path(output_text).expanduser()
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(all_text))
        
        print(f"✓ Text extracted successfully")
        print(f"  Input:  {input_docx}")
        print(f"  Output: {output_text}")
        print(f"  Total paragraphs: {len(doc.paragraphs)}")
        print(f"  Total tables: {len(doc.tables)}")
        
    except FileNotFoundError:
        print(f"Error: Input file not found: {input_docx}")
        sys.exit(1)
    except Exception as e:
        print(f"Error extracting text: {str(e)}")
        sys.exit(1)


def generate_template(input_docx: str, fields_json: str, output_docx: str) -> None:
    """
    Generate a template document with placeholders based on identified fields.
    
    Args:
        input_docx: Path to original .docx file
        fields_json: Path to fields.json (from LLM identification)
        output_docx: Path to save generated template
    """
    try:
        # Load original document
        doc = Document(input_docx)
        
        # Load fields mapping
        with open(Path(fields_json).expanduser(), 'r', encoding='utf-8') as f:
            fields_data = json.load(f)
        
        fields = fields_data.get('fields', [])
        if not fields:
            print("Warning: No fields found in fields.json")
            return
        
        # Create mapping: fieldName -> standardPath
        field_mapping = {
            field['fieldName']: field['standardPath']
            for field in fields
        }
        
        print(f"Field mapping loaded: {len(field_mapping)} fields")
        for name, path in field_mapping.items():
            print(f"  {name} → {{{{{path}}}}}")
        
        # Process paragraphs
        replaced_count = 0
        for para in doc.paragraphs:
            original_text = para.text
            modified_text = original_text
            
            # Search for field names and insert placeholders
            for field_name, standard_path in field_mapping.items():
                # Match patterns like "姓名：", "姓名:", "姓名 ", "姓名："
                patterns = [
                    f"{field_name}：",
                    f"{field_name}:",
                    f"{field_name} ",
                    f"{field_name}　",  # Full-width space
                ]
                
                for pattern in patterns:
                    if pattern in modified_text:
                        placeholder = f"{{{{{standard_path}}}}}"
                        # Insert placeholder after the field name
                        modified_text = modified_text.replace(
                            pattern,
                            f"{pattern}{placeholder}"
                        )
                        replaced_count += 1
                        print(f"  ✓ Inserted {{{{ {standard_path} }}}} after '{field_name}'")
            
            # Update paragraph text if modified
            if modified_text != original_text:
                para.text = modified_text
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    modified_text = original_text
                    
                    for field_name, standard_path in field_mapping.items():
                        patterns = [
                            f"{field_name}：",
                            f"{field_name}:",
                            f"{field_name} ",
                            f"{field_name}　",
                        ]
                        
                        for pattern in patterns:
                            if pattern in modified_text:
                                placeholder = f"{{{{{standard_path}}}}}"
                                modified_text = modified_text.replace(
                                    pattern,
                                    f"{pattern}{placeholder}"
                                )
                                replaced_count += 1
                                print(f"  ✓ Inserted {{{{ {standard_path} }}}} in table cell")
                    
                    if modified_text != original_text:
                        cell.text = modified_text
        
        # Save template
        output_path = Path(output_docx).expanduser()
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        print(f"\n✓ Template generated successfully")
        print(f"  Input:  {input_docx}")
        print(f"  Fields: {fields_json}")
        print(f"  Output: {output_docx}")
        print(f"  Placeholders inserted: {replaced_count}")
        
        if replaced_count == 0:
            print("\n⚠ Warning: No placeholders were inserted.")
            print("  Possible reasons:")
            print("  - Field names in document don't match fields.json")
            print("  - Field names have different punctuation or spacing")
            print("  - Fields are in images or special formats")
            print("\n  Suggested actions:")
            print("  - Review fields.json and document content")
            print("  - Manually adjust field names in fields.json")
            print("  - Verify document format is standard .docx")
        
    except FileNotFoundError as e:
        print(f"Error: File not found: {e.filename}")
        sys.exit(1)
    except json.JSONDecodeError:
        print(f"Error: Invalid JSON in fields file: {fields_json}")
        sys.exit(1)
    except Exception as e:
        print(f"Error generating template: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


def main():
    parser = argparse.ArgumentParser(
        description="Template Generator: Extract text and generate templates from Word documents"
    )
    
    parser.add_argument(
        '--action',
        choices=['extract', 'generate'],
        required=True,
        help="Action to perform: 'extract' text or 'generate' template"
    )
    
    parser.add_argument(
        '--input',
        required=True,
        help="Path to input .docx file"
    )
    
    parser.add_argument(
        '--output',
        required=True,
        help="Path to output file (text file for 'extract', .docx for 'generate')"
    )
    
    parser.add_argument(
        '--fields',
        help="Path to fields.json (required for 'generate' action)"
    )
    
    args = parser.parse_args()
    
    # Validate arguments
    if args.action == 'generate' and not args.fields:
        print("Error: --fields is required for 'generate' action")
        parser.print_help()
        sys.exit(1)
    
    # Expand ~ to home directory
    input_path = str(Path(args.input).expanduser())
    output_path = str(Path(args.output).expanduser())
    
    # Execute action
    if args.action == 'extract':
        extract_text(input_path, output_path)
    
    elif args.action == 'generate':
        fields_path = str(Path(args.fields).expanduser())
        generate_template(input_path, fields_path, output_path)


if __name__ == '__main__':
    main()
